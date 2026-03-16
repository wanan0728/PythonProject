"""
多模态文档处理器
功能：处理 DOCX、PDF、图片，提取文本和图片描述
"""
import os
import tempfile
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import hashlib
from datetime import datetime

from PIL import Image
import pytesseract
from docx import Document
from pdf2image import convert_from_path
import pymupdf4llm

from utils.logger import logger

# ===== 强制指定 Tesseract 路径（关键！）=====
TESSERACT_PATH = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
if os.path.exists(TESSERACT_PATH):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH
    print(f"✅ Tesseract 路径已设置: {TESSERACT_PATH}")

    # 测试 Tesseract 是否可用
    try:
        version = pytesseract.get_tesseract_version()
        print(f"✅ Tesseract 版本: {version}")

        # 查看可用的语言包
        langs = pytesseract.get_languages()
        print(f"✅ 可用语言包: {langs}")
    except Exception as e:
        print(f"❌ Tesseract 测试失败: {e}")
else:
    print(f"❌ Tesseract 未找到: {TESSERACT_PATH}")
    print("请从 https://github.com/UB-Mannheim/tesseract/wiki 下载安装")

# ========== 配置 ==========
TEMP_DIR = Path("temp/multimodal")
# 确保临时目录存在
TEMP_DIR.mkdir(parents=True, exist_ok=True)
print(f"📁 临时目录路径: {TEMP_DIR.absolute()}")
print(f"📁 目录是否存在: {TEMP_DIR.exists()}")


class MultimodalProcessor:
    """
    多模态文档处理器
    支持：DOCX（含图片）、PDF、图片
    """

    def __init__(self, use_image_captioning: bool = False):
        """
        初始化多模态处理器
        :param use_image_captioning: 是否使用AI生成图片描述（需要多模态模型）
        """
        self.use_image_captioning = use_image_captioning
        self.caption_model = None

        # 再次确认 Tesseract 可用
        try:
            # 创建一个简单的测试图片
            test_img = Image.new('RGB', (100, 30), color='white')
            from PIL import ImageDraw
            draw = ImageDraw.Draw(test_img)
            draw.text((10, 10), "Test", fill='black')
            test_text = pytesseract.image_to_string(test_img)
            print(f"✅ Tesseract 基础测试通过，识别结果: {test_text.strip() or '空'}")
        except Exception as e:
            print(f"⚠️ Tesseract 测试失败: {e}")

        if use_image_captioning:
            self._init_caption_model()

        logger.info("✅ 多模态处理器初始化完成")

    def _init_caption_model(self):
        """初始化图片描述模型（可选）"""
        try:
            from transformers import BlipProcessor, BlipForConditionalGeneration

            logger.info("加载图片描述模型 BLIP...")
            self.caption_processor = BlipProcessor.from_pretrained(
                "Salesforce/blip-image-captioning-base"
            )
            self.caption_model = BlipForConditionalGeneration.from_pretrained(
                "Salesforce/blip-image-captioning-base"
            )
            logger.info("✅ 图片描述模型加载完成")
        except Exception as e:
            logger.error(f"图片描述模型加载失败: {e}")
            self.use_image_captioning = False

    def process_file(self, file_path: Union[str, Path]) -> List[Dict[str, Any]]:
        """
        处理文件主入口
        :param file_path: 文件路径
        :return: 文档块列表，每块包含内容和元数据
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_ext = file_path.suffix.lower()
        logger.info(f"处理文件: {file_path.name} (类型: {file_ext})")

        if file_ext == '.docx':
            return self._process_docx(file_path)
        elif file_ext == '.pdf':
            return self._process_pdf(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            return self._process_image(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_ext}")

    def _process_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        """处理 DOCX 文件，提取文本和图片"""
        chunks = []
        doc = Document(file_path)

        # 提取所有段落
        full_text = []
        image_count = 0

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        # 提取图片
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_count += 1
                    image_data = rel.target_part.blob

                    # 确保目录存在
                    TEMP_DIR.mkdir(parents=True, exist_ok=True)

                    # 保存图片到临时文件
                    img_filename = f"{file_path.stem}_img_{image_count}.png"
                    img_path = TEMP_DIR / img_filename

                    with open(img_path, 'wb') as f:
                        f.write(image_data)

                    print(f"✅ 图片已保存: {img_path}")

                    # 对图片进行 OCR 或生成描述
                    image_text = self._process_image_file(img_path)

                    chunks.append({
                        "type": "image",
                        "content": image_text,
                        "metadata": {
                            "source": str(file_path),
                            "image_index": image_count,
                            "image_path": str(img_path),
                            "create_time": datetime.now().isoformat()
                        }
                    })

                except Exception as e:
                    logger.error(f"提取图片失败: {e}")
                    print(f"❌ 图片处理失败: {e}")

        # 添加文本块
        if full_text:
            chunks.append({
                "type": "text",
                "content": "\n".join(full_text),
                "metadata": {
                    "source": str(file_path),
                    "create_time": datetime.now().isoformat()
                }
            })

        logger.info(f"DOCX处理完成: {len(chunks)} 个块 (文本+{image_count}张图片)")
        return chunks

    def _process_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """处理 PDF 文件，使用 PyMuPDF4LLM"""
        chunks = []

        try:
            # 尝试用 PyMuPDF4LLM 提取文本和表格
            md_text = pymupdf4llm.to_markdown(str(file_path))

            if md_text and len(md_text.strip()) > 100:
                # 有足够文本，说明是纯文本 PDF
                chunks.append({
                    "type": "text",
                    "content": md_text,
                    "metadata": {
                        "source": str(file_path),
                        "parser": "pymupdf4llm",
                        "create_time": datetime.now().isoformat()
                    }
                })
                logger.info(f"PDF纯文本解析完成，长度: {len(md_text)} 字符")
            else:
                # 可能是扫描件，用 OCR
                logger.info("PDF可能是扫描件，使用OCR...")
                ocr_chunks = self._process_scanned_pdf(file_path)
                chunks.extend(ocr_chunks)

        except Exception as e:
            logger.error(f"PDF解析失败，尝试OCR: {e}")
            ocr_chunks = self._process_scanned_pdf(file_path)
            chunks.extend(ocr_chunks)

        return chunks

    def _process_scanned_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """处理扫描版 PDF（OCR）"""
        chunks = []

        try:
            # PDF 转图片
            images = convert_from_path(file_path, dpi=200)
            logger.info(f"PDF共 {len(images)} 页，开始OCR...")

            for i, img in enumerate(images, 1):
                # 保存临时图片
                img_path = TEMP_DIR / f"{file_path.stem}_page_{i}.png"
                img.save(img_path, 'PNG')

                # OCR
                text = self._ocr_image(img)

                if text.strip():
                    chunks.append({
                        "type": "text",
                        "content": text,
                        "metadata": {
                            "source": str(file_path),
                            "page": i,
                            "parser": "ocr",
                            "create_time": datetime.now().isoformat()
                        }
                    })

            logger.info(f"OCR完成，生成 {len(chunks)} 个文本块")

        except Exception as e:
            logger.error(f"OCR处理失败: {e}")

        return chunks

    def _process_image(self, file_path: Path) -> List[Dict[str, Any]]:
        """处理单独图片文件"""
        return [{
            "type": "image",
            "content": self._process_image_file(file_path),
            "metadata": {
                "source": str(file_path),
                "image_path": str(file_path),
                "create_time": datetime.now().isoformat()
            }
        }]

    def _ocr_image(self, img: Image.Image) -> str:
        """对图片进行OCR识别（多种语言尝试）"""
        # 尝试多种语言组合
        lang_configs = [
            'chi_sim+eng',  # 简体中文+英文
            'eng',          # 纯英文
            'chi_sim',      # 纯中文
            'eng+chi_sim',  # 英文+中文（顺序不同）
        ]

        for lang in lang_configs:
            try:
                text = pytesseract.image_to_string(img, lang=lang)
                if text and len(text.strip()) > 10:  # 有足够多的文字
                    logger.debug(f"OCR成功，语言: {lang}，长度: {len(text)}")
                    return text.strip()
            except Exception as e:
                logger.debug(f"OCR尝试失败 {lang}: {e}")
                continue

        # 所有尝试都失败，返回空
        return ""

    def _process_image_file(self, img_path: Path) -> str:
        """处理单张图片，返回文字描述"""
        try:
            # 打开图片
            img = Image.open(img_path)

            # 尝试 OCR 识别
            ocr_text = self._ocr_image(img)

            if ocr_text:
                logger.info(f"✅ 图片OCR识别成功，文字长度: {len(ocr_text)}")
                return f"[图片中的文字: {ocr_text}]"
            else:
                logger.warning(f"⚠️ 图片未识别出文字: {img_path}")

                # 如果开启了图片描述功能，尝试生成描述
                if self.use_image_captioning and self.caption_model:
                    caption = self._generate_caption(img)
                    if caption:
                        return f"[图片描述: {caption}]"

                return "[图片内容无法识别]"

        except Exception as e:
            logger.error(f"图片处理失败: {e}")
            return "[图片内容无法识别]"

    def _generate_caption(self, img: Image) -> str:
        """生成图片描述"""
        try:
            inputs = self.caption_processor(img, return_tensors="pt")
            out = self.caption_model.generate(**inputs)
            caption = self.caption_processor.decode(out[0], skip_special_tokens=True)
            return caption
        except Exception as e:
            logger.error(f"图片描述生成失败: {e}")
            return ""


# ========== 文档加载器（兼容 LangChain）==========
from langchain_core.documents import Document as LCDocument

class MultimodalLoader:
    """LangChain 兼容的多模态文档加载器"""

    def __init__(self, processor: MultimodalProcessor):
        self.processor = processor

    def load(self, file_path: Union[str, Path]) -> List[LCDocument]:
        """加载文件并转换为 LangChain Document 列表"""
        chunks = self.processor.process_file(file_path)

        docs = []
        for chunk in chunks:
            doc = LCDocument(
                page_content=chunk["content"],
                metadata=chunk["metadata"]
            )
            docs.append(doc)

        return docs